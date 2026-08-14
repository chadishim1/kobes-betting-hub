from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT="docs/Instagram_Marketing_Operations_Package.docx"
NAVY="14283D"; BLUE="2E74B5"; PALE="EAF2F8"; RED="9B1C1C"; WHITE="FFFFFF"; GRAY="5B6573"
d=Document(); s=d.sections[0]
s.top_margin=Inches(.72); s.bottom_margin=Inches(.68); s.left_margin=Inches(.78); s.right_margin=Inches(.78)

def font(run,size=10,bold=False,color=NAVY,italic=False):
 run.font.name="Aptos"; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"),"Aptos"); run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"),"Aptos"); run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=RGBColor.from_string(color)
st=d.styles["Normal"]; st.font.name="Aptos"; st.font.size=Pt(9.5); st.paragraph_format.space_after=Pt(4); st.paragraph_format.line_spacing=1.08
for n,z,c,b,a in [("Heading 1",16,BLUE,14,6),("Heading 2",12,NAVY,10,4),("Heading 3",10.5,GRAY,7,3)]:
 x=d.styles[n]; x.font.name="Aptos Display"; x.font.size=Pt(z); x.font.bold=True; x.font.color.rgb=RGBColor.from_string(c); x.paragraph_format.space_before=Pt(b); x.paragraph_format.space_after=Pt(a); x.paragraph_format.keep_with_next=True

def p(t="",size=9.5,bold=False,color=NAVY,align=None,italic=False,after=4):
 q=d.add_paragraph(); q.paragraph_format.space_after=Pt(after); q.alignment=align; font(q.add_run(t),size,bold,color,italic); return q
def bullet(t):
 q=d.add_paragraph(style="List Bullet"); q.paragraph_format.space_after=Pt(3); q.paragraph_format.left_indent=Inches(.25); q.paragraph_format.first_line_indent=Inches(-.15); q.add_run(t)
def number(t):
 q=d.add_paragraph(style="List Number"); q.paragraph_format.space_after=Pt(3); q.paragraph_format.left_indent=Inches(.25); q.paragraph_format.first_line_indent=Inches(-.15); q.add_run(t)
def shade(c,fill):
 sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),fill); c._tc.get_or_add_tcPr().append(sh)
def table(rows,widths,fs=7.6):
 t=d.add_table(rows=0,cols=len(widths)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
 for ri,row in enumerate(rows):
  cells=t.add_row().cells
  for i,v in enumerate(row):
   cells[i].width=Inches(widths[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
   if ri==0: shade(cells[i],NAVY)
   q=cells[i].paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(str(v)),fs,ri==0,WHITE if ri==0 else NAVY)
 t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader")); d.add_paragraph().paragraph_format.space_after=Pt(1); return t
def box(label,text,red=False):
 t=d.add_table(rows=1,cols=1); t.autofit=False; c=t.cell(0,0); shade(c,"FCE8E6" if red else PALE); q=c.paragraphs[0]; q.paragraph_format.space_after=Pt(0); font(q.add_run(label+"  "),9,True,RED if red else NAVY); font(q.add_run(text),9,False,RED if red else NAVY); d.add_paragraph().paragraph_format.space_after=Pt(1)

h=s.header.paragraphs[0]; h.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font(h.add_run("KOBE'S BETTING HUB / INSTAGRAM OPS"),7.5,True,GRAY)
f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER; font(f.add_run("Internal playbook • v1.0 • 13 August 2026"),7.5,False,GRAY)

p("OPERATIONS PLAYBOOK",10,True,"C18B2E",WD_ALIGN_PARAGRAPH.CENTER,after=15)
p("Instagram Marketing\nOperations Package",27,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=7)
p("Daily Reels, feed posts and Stories—with a review-first source and approval system",13,False,BLUE,WD_ALIGN_PARAGRAPH.CENTER,after=24)
box("NON-NEGOTIABLE", "Use owned, commissioned, licensed, or explicitly approved assets only. Do not access partner Discord communities. Do not repost unapproved third-party media.",True)
p("Built for Kobe's Betting Hub",13,True,NAVY,WD_ALIGN_PARAGRAPH.CENTER,after=18)
p("Purpose: create a consistent, measurable Instagram presence without introducing rights, claims, account-security, or automated-publishing risk.",10,False,GRAY,WD_ALIGN_PARAGRAPH.CENTER,italic=True)

d.add_page_break(); d.add_heading("1. Operating system",1)
p("The weekly plan supplies one daily Reel, one daily feed post, and one Story sequence. Templates accelerate drafting; humans approve rights, facts, claims, and the exact final version.")
box("DAILY OUTPUT", "1 Reel + 1 feed post + 3–6 Story frames. If a pick, result, line, offer, or asset is not approved in time, publish an evergreen educational or community post.")
d.add_heading("Roles and controls",2)
table([["Role","Owns","Hard gate"],["Content lead","Calendar, briefs, copy, queue","No draft without source/asset record"],["Designer/editor","Creative and exports","No unapproved asset or altered claim"],["Kobe/admin","Facts, tone, offer, final approval","Required for betting-specific post"],["Publisher","Schedule/post, URL and status log","Exact approved version only"],["Community owner","Replies, escalation, insight log","No sensitive data or individualized stakes advice"]],[1.05,3.05,2.75])
for x in ["Intake by 8:00 a.m.: approved picks/results, owned footage, offers, deadlines and restrictions.","Draft by 10:00 a.m.: all three surfaces, caption, CTA, alt/accessibility treatment and source IDs.","Review by noon: verify event, time zone, line/odds timestamp, units, record math, permission, disclosures and links.","Approve and schedule: approver signs the exact export/caption version; publisher uses Meta Business Suite or manual Instagram first.","Closeout: log URLs, time, reach, watch time, saves, shares, profile/link actions, replies and lessons."]: number(x)
d.add_heading("Green / yellow / red",2)
table([["Zone","Examples","Action"],["GREEN","Original footage; owned graphics; approved pick/result data; general education","Normal review queue"],["YELLOW","Partner quote/logo; screenshot; music; testimonial; customer message","Hold for written permission and scope"],["RED","Private/Discord content; unapproved media; false wins; guarantees; minors; personal data","Do not use"]],[.8,4.0,2.05])

d.add_page_break(); d.add_heading("2. Seven-day daily calendar",1)
p("Repeat this spine weekly. Posting windows are hypotheses; adjust after four weeks of Insights data.")
rows=[["Day / pillar","Reel","Feed","Stories","Goal / CTA"],
["Mon / Proof","Week reset + honest prior-week lesson","Weekly slate / release windows","Recap → schedule → question box","Trust / save"],
["Tue / Teach","20-sec market or line lesson","3-point educational carousel","Quiz → answer → takeaway","Saves / share"],
["Wed / Process","Behind-the-scenes approval flow","How a pick becomes approved","Desk clip → checklist → AMA","Trust / ask"],
["Thu / Community","Kobe answers one FAQ","One approved member benefit","Question sticker → answers","Replies / “HUB”"],
["Fri / Slate","Three weekend watch items","Matchup/watch-list card","Countdown → poll → reminder","Reach / follow"],
["Sat / Game day","Original analysis or setup","Approved pick or evergreen fallback","Status → context → update","Profile / bio"],
["Sun / Recap","Verified results + one lesson","W–L/units recap + method note","Results → lesson → next week","Trust / notify"]]
table(rows,[.75,1.65,1.65,1.65,1.1],6.9)
d.add_heading("Daily timing",2)
table([["Window","Surface","Job"],["8–10 a.m.","Story opener/poll","Signal activity and collect input"],["11 a.m.–1 p.m.","Feed","Saveable detail and profile credibility"],["3–6 p.m.","Reel","Discovery and event-adjacent attention"],["Post-event","Story update/recap","Close loop; remove or correct stale claims"]],[1.25,1.75,3.8])
box("EVERGREEN FALLBACK", "Betting terms, record-keeping, review process, responsible play, a community FAQ, or an original founder clip. Never rush an unverified result or stale line.")

d.add_page_break(); d.add_heading("3. Production briefs",1)
briefs=[
("Monday — Reset & proof","Hook: “New week. Same rule: show the work.”","Reel: Kobe on camera → verified prior-week result → this week’s schedule.","Feed: weekly slate with expected release windows and “subject to approval.”","Stories: result context → schedule → question box.","Proof: approved recap, units convention, date range, source IDs."),
("Tuesday — Teach","Hook: “The line moved. Here’s what that actually tells us.”","Teach one concept, one original/licensed example and one takeaway; no guaranteed edge.","Feed: definition → why it matters → mistake → checklist → CTA.","Stories: quiz → answer → save/share prompt.","Proof: public authoritative or licensed data, observed timestamp."),
("Wednesday — Process","Hook: “What happens before a play ever gets posted?”","Show intake → verify → approve → publish → grade → recap.","Feed: five-stage process carousel.","Stories: sanitized owned workspace clips; never private channels, member data or tokens.","Proof: sanitized screens and owner approval."),
("Thursday — Community","Hook: “The question I get every week…”","Answer one membership/process FAQ directly.","Feed: one approved benefit without outcome promises.","Stories: question sticker → answers → support route.","Proof: approved offer, price, eligibility and support link."),
("Friday — Watch list","Hook: “Three things I’m watching this weekend.”","Three items, clearly labeled watch list—not final plays.","Feed: matchup grid or one deep dive.","Stories: countdown and poll; update stale information.","Proof: schedules from authoritative public sources; no broadcast imagery."),
("Saturday — Game day","Hook: “Here’s the setup—not a victory lap.”","Original face-to-camera analysis plus owned graphics.","Feed: approved pick only; otherwise evergreen checklist.","Stories: release status → details → result update.","Proof: line/odds/time/source and grading rule."),
("Sunday — Recap","Hook: “The number is the number. Here’s what we learned.”","Verified record first; one win, one miss, one adjustment.","Feed: W–L, units and defined date range.","Stories: transparent results → lesson → next-week prompt.","Proof: reconciled pick log and sign-off.")]
for title,*items in briefs:
 d.add_heading(title,2)
 for x in items: bullet(x)

d.add_page_break(); d.add_heading("4. Source and asset workflow",1)
for x in ["Register: create one source/asset record and unique ID before drafting.","Classify: owner, origin, reuse permission, credit, platforms, organic/paid scope, territory and expiration.","Verify: record fact URL or approved internal record, observed timestamp/time zone, event, market, line/odds and units convention.","Clear: mark Green only with documentation. Publicly visible does not mean reusable.","Produce: work from approved IDs; edits cannot change a pick/result or hide attribution.","Approve: sign final exported file and caption together, with version and timestamp.","Archive: retain final asset, caption, approval, URL, results and expiry/takedown trigger."]: number(x)
d.add_heading("Required source/asset record",2)
table([["Field","Record"],["Asset ID/version","Unique ID + v1/v2/FINAL"],["Origin/creator","Owned, commissioned, licensed, partner-supplied, platform library"],["Permission evidence","Contract, license, written approval or ownership record"],["Usage scope","Organic/paid; platform/placement; territory; start/end"],["Credit/restrictions","Exact credit, edit, likeness and music limits"],["Fact source","URL or approved internal record + observed timestamp"],["Approvals/archive","Rights, facts, brand, publisher; final path and post URL"]],[1.65,5.2])
box("DISCORD BOUNDARY", "Do not enter, scrape, screenshot, monitor, summarize, or republish partner Discord content. A partner may separately supply an asset through an approved channel with explicit written usage rights.",True)

d.add_page_break(); d.add_heading("5. Creative standards",1)
d.add_heading("Master brief fields",2)
for x in ["Objective and one audience state","One message the viewer should remember","Approved proof and source/asset IDs","Hook and beat-by-beat structure","Format, ratio, duration/slides, cover and caption","One CTA and tested destination","Rights, claim, disclosure, paid-use and jurisdiction check","Named approvers, deadline and primary metric"]: bullet(x)
d.add_heading("Reels",2)
for x in ["9:16 vertical; keep essential copy clear of interface edges and test profile-grid cover crop.","0–2s hook; 2–8s context; 8–20s proof; final 2–4s CTA. One idea per Reel.","Use original voice or properly licensed audio. Correct auto-captions, every number and every name.","Deliver clean master, captioned master, cover, caption, accessibility note, asset IDs and approvals."]: bullet(x)
d.add_heading("Feed",2)
for x in ["Use 4:5 portrait where practical. Slide 1 has one short promise; later slides carry proof.","Carousel: hook → context → evidence → takeaway → CTA. Recaps state date range and methodology.","No faux sportsbook UI, unlicensed team/broadcast photography or copied third-party design."]: bullet(x)
d.add_heading("Stories",2)
for x in ["9:16, one idea per frame, large type, 3–6 frames; stickers must not cover proof/disclosure.","Sequence: attention → context → proof → interaction → CTA. Label timestamps and time zone.","Reuse approved evergreen sequences only; never auto-recycle time-sensitive picks, lines, offers or results."]: bullet(x)

d.add_page_break(); d.add_heading("6. Caption and CTA library",1)
p("Voice assumption: direct, accountable, useful and low-hype. Replace after Kobe’s writing samples are approved.")
cap=[["Use","Starter","CTA"],
["Reset","New week. Here’s what we’re watching, when updates are expected, and how the process works. Nothing gets posted just to fill the board.","Save this schedule."],
["Teach","Quick breakdown: [TERM]. It matters because [REASON]. The mistake is treating one signal like a guarantee.","Save/share."],
["Process","Before a play goes live: source checked, market verified, details reviewed, then approval. The work comes before the post.","What next?"],
["Watch list","On the radar: [A], [B], [C]. Watch list—not final plays. Releases include the current line and timestamp.","Follow."],
["Pick","[SPORT] — [EVENT]. Play: [PICK]. Line/odds observed [TIME/TZ]. Units: [X]. [WHY]. Lines move; confirm before acting.","Link in bio."],
["Recap","[RANGE]: [W]–[L], [UNITS]. Right: [X]. Missed: [Y]. Next adjustment: [Z].","Share."],
["No play","No forced action today. Passing is part of the process when the number or information is not there.","Notify."],
["Membership","Betting Hub includes [BENEFITS]. Price: [PRICE/TERM]. No promises—just process, updates and transparent recaps.","Read details."],
["Responsible","Set limits before the slate. Never chase, never treat a pick as guaranteed, and step away when betting stops being fun.","Save."]]
table(cap,[1.0,4.9,.85],6.9)
d.add_heading("CTA bank",2)
for x in ["Save this breakdown.","Share it with someone building a process.","Follow for the verified update.","Turn on notifications for release windows.","Reply with the topic you want next.","Use the link in bio for approved membership details.","Send “HUB” for the public information link.","Read the full terms before joining."]: bullet(x)
box("CTA RULE", "Use one primary CTA. No guaranteed wins, easy money, fear-based urgency, inside-information implication, or untested destination.")

d.add_page_break(); d.add_heading("7. Approval, publishing and correction",1)
d.add_heading("Final preflight",2)
for x in ["Permission evidence covers this platform, placement and paid/organic use.","No Discord/private content, unapproved media, personal data or member identifiers.","Event, market, pick, line/odds, units, timestamp/time zone, record and math match the approved record.","No guaranteed-win, risk-free, easy-money, misleading-performance or fabricated-scarcity claim.","Caption, on-screen text, spoken audio, cover, tags, link and disclosure reviewed together.","Publishing destination and cross-post settings are correct.","Paid promotion has separate compliance approval, Meta eligibility/permission and lawful age/location controls.","Kobe/admin and publisher approved the exact final version.","Post URL, monitoring owner and correction/takedown route are assigned."]: bullet("☐ "+x)
d.add_heading("Correction protocol",2)
for x in ["Pause scheduling/promotion.","Capture the live post and document the issue.","Notify approver; choose edit, clarification, deletion or takedown.","Clearly correct any error that could affect a betting decision or performance claim.","Update the record, root cause and checklist; retain the audit trail."]: number(x)
d.add_heading("Comment and DM rules",2)
for x in ["Answer public process/offer questions with approved copy; move account-specific support to the secure route.","Never request passwords, full payment-card data or sensitive identity data in DMs.","Do not personalize stakes or advise chasing. Escalate minors, fraud, impersonation, threats, crisis and payment disputes."]: bullet(x)

d.add_page_break(); d.add_heading("8. Professional account and Meta integration",1)
p("Reconfirm current Meta rules immediately before setup or API development; features and permission names change.")
d.add_heading("Baseline",2)
for x in ["Use an Instagram professional account (Business or Creator). Professional accounts are public; Business is the cleaner default for a paid service.","Connect the official Facebook Page. The connecting person needs appropriate Page/Facebook access.","Use the correct Meta business portfolio, named individual access, least privilege, 2FA and at least two trusted admins.","Record owners for Page, Instagram, ad account, domain, payment and recovery. Never share passwords/tokens in documents or chat.","Start with Meta Business Suite/manual Instagram. Test publishing destination, cross-posting, inbox permissions and Insights."]: bullet(x)
d.add_heading("API / third-party readiness",2)
table([["Requirement","Implication"],["Supported account/path","Confirm chosen Instagram API login path supports account and surface"],["Developer app","Business ownership; privacy/data deletion; app review/verification where required"],["Permissions","Request only current account/media/publishing access needed"],["Tokens","Secrets manager, least scope, rotation/revocation, access logging; never commit"],["Media/specs","Reachable media URL; validate current format, ratio, duration and size"],["Publish controls","Approve → validate → container → poll → publish → log; idempotent retry"],["Feature eligibility","Do not promise Story/Reel/carousel publishing until tested"],["Monitoring","Alert on token, permission, processing, rate-limit and destination failures"]],[1.65,5.2],7.2)
d.add_heading("Betting-content gate",2)
box("PRIOR META PERMISSION MAY BE REQUIRED", "Meta’s published community guidance says accounts promoting online gambling, online real-money games of skill, or online lotteries must obtain prior written permission. Treat promotion of a paid betting-picks membership as legal/compliance review territory; organic eligibility does not equal ad eligibility.",True)
for x in ["Obtain qualified legal/compliance review for each targeted jurisdiction and the exact service/claims.","Before paid campaigns, verify Meta authorization, applicable licenses/partner status, landing-page compliance and lawful age/location restrictions.","Keep organic and paid approvals separate; boosting is a new use.","Never target minors or imply certainty, recovery of losses or financial security."]: bullet(x)

d.add_page_break(); d.add_heading("9. Measurement and 30-day rollout",1)
table([["Surface","Primary","Supporting","Decision"],["Reels","Retention/watch time","Completion, shares, follows, profile visits","Hook/length"],["Feed","Saves per reach","Shares, comments, profile visits","Topic/design"],["Stories","Completion","Exits, replies, stickers, link taps","Sequence/CTA"],["Membership","Qualified page action","Link taps, starts, joins, support","Offer/message"],["Operations","On-time approved rate","Rework, rights holds, corrections","Workflow"]],[1.0,1.65,2.5,1.7],7.1)
d.add_heading("Weekly review",2)
for x in ["Pull Insights and log for the same date range.","Name top two posts by primary metric—not likes alone.","Choose one hook, format and CTA to repeat; one to stop; one to test.","Review rights holds, corrections, support issues and delays.","Lock next week’s pillars, owners and evidence needs."]: number(x)
d.add_heading("30-day rollout",2)
table([["Week","Focus","Exit criterion"],["1 Foundation","Permissions, templates, account setup; manual publish","7 days with complete approval records"],["2 Consistency","Full cadence; log time and holds","≥85% approved slots on time"],["3 Learning","Repeat pillars; test one hook and CTA","Useful benchmarks established"],["4 Control","Audit rights, claims, access, corrections, results","Only stable low-risk automation selected"]],[1.05,3.5,2.3])
d.add_heading("Inputs still required from Kobe",2)
for x in ["20–50 voice samples and phrases to use/avoid","Membership promise, price, terms and returning-member offer","Owned brand assets and visual rules","Named approvers, publishers, support and compliance owner","Approved source list, attribution rules and rights evidence","Units definition, grading and record-display method","Target jurisdictions"]: bullet("☐ "+x)

d.add_page_break(); d.add_heading("Appendix A. Daily assignment",1)
for x in ["Date / publish window","Pillar / objective","Audience state","Reel hook + beats","Feed concept + slide plan","Story sequence","Proof/source IDs","Asset IDs/rights status","Caption/CTA/destination","Disclosure","Designer/editor","Fact reviewer","Kobe/admin approval","Publisher","Primary metric","Final URLs/results/learning"]: p(x+":  ______________________________________________________________",9,True,NAVY,after=7)
d.add_heading("Appendix B. Official Meta references",1)
p("Checked 13 August 2026; recheck before implementation.",8.5,False,GRAY,italic=True)
for x in ["Professional accounts: https://www.facebook.com/help/instagram/138925576505882","Connect professional Instagram and Page: https://www.facebook.com/help/instagram/402748553849926","Connect/disconnect Instagram and Page: https://www.facebook.com/help/1148909221857370","Effects of connecting accounts: https://www.facebook.com/help/2546917405323366","Meta Commercial Terms: https://www.facebook.com/legal/commercial_terms","Community guidance / gambling permission: https://www.facebook.com/help/477434105621119","Audience age/location controls: https://www.facebook.com/help/157306091096340"]: bullet(x)
p("This is an operating playbook, not legal advice. Availability and requirements vary by account, country, product and API path.",9,True,RED)

d.core_properties.title="Instagram Marketing Operations Package"; d.core_properties.author="Kobe's Betting Hub"
d.save(OUT); print(OUT)
